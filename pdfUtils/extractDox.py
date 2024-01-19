import json
import os.path
from time import sleep

import convertapi
from docx import Document
from fitz import fitz
from loguru import logger

def extract_text(dox_dir: str):

    text = {
        "Title": "",
        "Navigation Bra": "null",
        "AuthorInfo": "",
        "Abstract": "",
        "Content": [],
    }
    textContent = {
        "Title": "",
        "Content": "",
        "Below": [],
        "Grade": 0,
    }
    storedText = []

    # 打开Word文档
    doc = Document(dox_dir)
    navigationBra = 1
    authorContent = 0
    paperDict = {}  # 保存的内容为'contextType' = 'content'
    pre_paragraph = doc.paragraphs[0]
    contentType2 = 'No classification'

    #取出目录
    pdf = fitz.open(dox_dir.replace('docx', 'pdf'))
    toc = pdf.get_toc()  # 获取pdf目录
    catalogue = []
    for t in toc:
        # if t[0] == 1:
        # print(t)
        catalogue.append(t[1].lower())
    # print(catalogue)

    abstract = 0 if 'abstract' in catalogue else 1

    # 遍历文档中的段落并打印文本内容
    test = 1
    for paragraph in doc.paragraphs:


        tmp = paragraph.text.replace(" ", "").replace("\n", "")
        if not tmp:
            continue
        # 获取段落的样式
        style = paragraph.style.name
        for run in paragraph.runs:
            # 获取文本的具体属性
            font_name = run.font.name  # 字体名称
            font_size = run.font.size  # 字体大小
            bold = run.font.bold  # 是否加粗
            italic = run.font.italic  # 是否斜体
            underline = run.font.underline  # 是否有下划线
            # 获取段落的对齐方式
            alignment = paragraph.alignment
            # 根据对齐方式输出文字对齐方式
            # if alignment == 0:
            #     #print("左对齐")
            # elif alignment == 1:
            #     #print("居中对齐")
            # elif alignment == 2:
            #     #print("右对齐")
            # elif alignment == 3:
            #     #print("两端对齐")
            # else:
            # print("未知对齐方式")
            break
        if style == 'Title':
            contentType = 'Title'
            content = paragraph.text
            navigationBra = 0
            authorContent = 1
            paperDict.update({contentType: content})
            text["Title"] = content

        elif navigationBra == 1:
            contentType = 'Navigation Bra'
            content = paragraph.text
            navigationBra = 0
            paperDict.update({contentType: content})
            text["Navigation Bra"] = content

        elif alignment == 1 and authorContent == 1:
            contentType = 'Authors Information'
            content = paragraph.text
            if not contentType in paperDict:
                paperDict[contentType] = paragraph.text
                text["AuthorInfo"] = content
            else:
                paperDict[contentType] += '\n' + paragraph.text
                text["AuthorInfo"] += '\n' + content


        elif alignment == 3 and abstract == 1:
            authorContent = 0
            contentType = 'Abstract'
            content = paragraph.text
            if not contentType in paperDict:
                paperDict[contentType] = paragraph.text
                text["Abstract"] = content
            else:
                paperDict[contentType] += '\n' + paragraph.text
                text["Abstract"] += '\n' + content

        else:
            authorContent = 0
            abstract = 0
            if paragraph.text.lower() in catalogue:
                contentType2 = paragraph.text

                if textContent["Title"]:
                    storedText.append({
                        "Title": textContent["Title"],
                        "Content": textContent["Content"],
                        "Below": [],
                        "Grade": textContent["Grade"],
                    })
                    # print(f'test here:\n{textContent["Title"]}\n\ntest over\n')
                textContent["Title"] = paragraph.text
                textContent["Content"] = ""
                change = 1
                for t in toc:
                    if paragraph.text.lower() == t[1].lower():
                        textContent["Grade"] = t[0]
                        break

            else:
                change = 0
            # else:
            #     logger.debug(f'paragraph.text = \n{paragraph.text}\ncatalogue2 = \n{catalogue}')
            if not contentType2 in paperDict:
                paperDict[contentType2] = paragraph.text
            else:
                paperDict[contentType2] += '\n' + paragraph.text

            if change == 0:
                textContent["Content"] += '\n' + paragraph.text

        test += 1
        # if test == 36:
        #     break

    storedText.append({
        "Title": textContent["Title"],
        "Content": textContent["Content"],
        "Below": [],
        "Grade": textContent["Grade"],
    })
    # print(f'test here:\n{textContent["Title"]}\n\ntest over\n')


    # print(f'test2 here:\n{len(storedText)}\n\ntest2 over\n')
    # print(storedText)
    num = len(storedText)
    for i in range(num):
        # print(storedText[i])
        # print(f'test2 here:\n{storedText[i]["Grade"]}\n{storedText[i]["Title"]}\ntest2 over\n')
        if storedText[i]["Grade"] == 1:
            res = getRelation(storedText, i)
            # print(f'test1 here:\n{res}\n\ntest1 over\n')
            # print(f'test3 here:\n{text}\n\ntest3 over\n')
            belowcontent = []
            for item in res["Below"]:
                belowcontent.append(item)
            text["Content"].append({
                "Title": res["Title"],
                "Content": res["Content"],
                "Below": belowcontent,
                "Grade": res["Grade"],
            })
            # print(f'test2 here:\n{text}\n\ntest2 over\n')


    # print(text)
    with open('work_file', 'w') as f:
        json.dump(text, f) #ensure_ascii=False
        logger.info("json文件写入完成...")
    return paperDict

def getRelation(storedText: [], i: int):
    for j in range(i + 1, len(storedText)):
        if storedText[j]["Grade"] > storedText[i]["Grade"] and storedText[j]["Grade"] == storedText[i]["Grade"] + 1:
            # print(f'test3 here:\n{storedText[i]["Grade"]}\n{storedText[i]["Title"]}\ntest3 over\n')
            # print(f'test4 here:\n{storedText[j]["Grade"]}\n{storedText[j]["Title"]}\ntest4 over\n')
            res = getRelation(storedText, j)
            belowcontent = []
            for item in res["Below"]:
                belowcontent.append(item)
            storedText[i]["Below"].append({
                "Title": res["Title"],
                "Content": res["Content"],
                "Below": belowcontent,
                "Grade": res["Grade"],
            })
        else:
            # print('pass\n')
            return storedText[i]
    return storedText[i]

def pdf_to_docx(PDFDir):
    convertapi.api_secret = 'UNQFiyvf08nxcmgH'
    convertapi.convert('docx', {
        'File': PDFDir
    }, from_format='pdf').save_files('./')

def extract_text_from_pdf(pdf_path):
    PDFDir = pdf_path
    DOCXDir = PDFDir.replace(".pdf", ".docx")
    #
    # 将pdf转为word，如果该word已经存在则不执行
    if not (os.path.exists(PDFDir.replace(".pdf", ".docx"))):
        logger.info('Convertion starts...')
        pdf_to_docx(PDFDir)
        # 移动“../tmp.docx”到"./tmp.docx"
        os.rename('./tmp.docx', './pdfUtils/tmp.docx')
    else:
        logger.info('file has been converted')

    # 调用输出解析内容的方法
    logger.info('Extraction starts...')
    mydict = extract_text(DOCXDir)
    logger.info('Extraction finishes...')

if __name__ == '__main__':
    extract_text_from_pdf('dong.pdf')
